#!/usr/bin/env python3
"""将 项目开发报告.md 转换为格式良好的 Word 文档 (.docx)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import os

INPUT_FILE = r'C:\Users\20822\music-recommendation\项目开发报告.md'
OUTPUT_FILE = r'C:\Users\20822\music-recommendation\项目开发报告.docx'

doc = Document()

# ========== 页面设置 ==========
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# ========== 样式定义 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题样式
for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading_font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_font.size = Pt(22)
    elif i == 2:
        heading_font.size = Pt(16)
    elif i == 3:
        heading_font.size = Pt(14)
    elif i == 4:
        heading_font.size = Pt(12)


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_paragraph(text, bold=False, font_size=None, alignment=None, font_name=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name or '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name or '宋体')
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if alignment is not None:
        p.alignment = alignment
    return p


def parse_and_convert(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    in_list = False

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                if code_lines:
                    for cl in code_lines:
                        p = doc.add_paragraph()
                        p.style = doc.styles['Normal']
                        run = p.add_run(cl)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                        pf = p.paragraph_format
                        pf.left_indent = Cm(0.5)
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)
                code_lines = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                # 检查是否有语言标识 inline? 不处理，跳过该行
                i += 1
                continue

        if in_code_block:
            code_lines.append(line.rstrip('\n'))
            i += 1
            continue

        # 表格
        if line.strip().startswith('|') and line.strip().endswith('|'):
            row_cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            # 跳过分隔行
            if all(re.match(r'^[-:]+$', c) for c in row_cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(row_cells)
            i += 1
            continue
        else:
            if in_table:
                # 结束表格，写入 Word 表格
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for ri, row in enumerate(table_rows):
                        for ci in range(num_cols):
                            cell = table.cell(ri, ci)
                            cell_text = row[ci] if ci < len(row) else ''
                            # 清空默认段落
                            cell.paragraphs[0].clear()
                            run = cell.paragraphs[0].add_run(cell_text)
                            run.font.name = '宋体'
                            run.font.size = Pt(9)
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                            if ri == 0:
                                # 表头
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                set_cell_shading(cell, '4472C4')
                            elif ri % 2 == 0:
                                set_cell_shading(cell, 'F2F2F2')

                    doc.add_paragraph()  # 表后空行
                table_rows = []
                in_table = False

        # 处理标题
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            doc.add_heading(title, level=level)
            i += 1
            continue

        # 水平线
        if line.strip() == '---':
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # 无序列表
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if list_match:
            text = list_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            # 清除默认文本
            p.clear()
            # 处理 bold
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue

        # 有序列表
        ordered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ordered_match:
            text = ordered_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue

        # 引用块
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # 普通段落
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph()
            # 处理行内代码 `code`
            parts = re.split(r'(`[^`]+`)', stripped)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
                else:
                    # 处理 bold **text**
                    bold_parts = re.split(r'(\*\*.*?\*\*)', part)
                    for bp in bold_parts:
                        if bp.startswith('**') and bp.endswith('**'):
                            run = p.add_run(bp[2:-2])
                            run.bold = True
                        else:
                            run = p.add_run(bp)
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        i += 1

    # 处理末尾可能的未关闭表格
    if in_table and table_rows:
        num_cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(table_rows):
            for ci in range(num_cols):
                cell = table.cell(ri, ci)
                cell_text = row[ci] if ci < len(row) else ''
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.name = '宋体'
                run.font.size = Pt(9)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_shading(cell, '4472C4')


print("Converting...")
parse_and_convert(INPUT_FILE)

doc.save(OUTPUT_FILE)
print(f"Done: {OUTPUT_FILE}")
print(f"File size: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")
